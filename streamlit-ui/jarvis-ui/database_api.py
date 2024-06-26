from docx import Document as docxDocument
from langchain.docstore.document import Document as LangchainDocument
from langchain_community.document_loaders.llmsherpa import LLMSherpaFileLoader
from langchain_community.vectorstores import Neo4jVector
from langchain_community.embeddings import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import os

config={"ollama_base_url": "http://ollama:11434",
        "llm_name": "llama3",
        "neo4j_url": "bolt://neo4j:7687",
        "neo4j_username": "neo4j",
        "neo4j_password": "password",		
        }

# load embedding model
embeddings = OllamaEmbeddings(
    base_url=config["ollama_base_url"],	
    model=config["llm_name"]
)

def docParser(file_path, st):
    """
    Parses a document file and returns a split up version of the document.
    Requires file type to be reflected in the file extension.

    Args:
        file_path (str): The path to the document file

    Returns:
        List of LangchainDocument objects
    """

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=750, chunk_overlap=20)
    file_name = os.path.basename(file_path)

    # This assumes that we can tell file type from file extension
    # May not work for linux based systems
    # TODO: NOT THE BEST CHUNKING STRATEGY
    if file_path.endswith('.docx'):
        doc_splits = []
        doc = docxDocument(file_path)

        para_num = 0
        for para in doc.paragraphs:
            if not para.text:
                continue

            langchain_doc_splits = LangchainDocument(
                page_content=para.text,
                metadata={
                    "source": file_path,
                    "chunk_number": para_num,
                    "chunk_type": "para"
                }
            )
            doc_splits.append(langchain_doc_splits)
            para_num += 1

        return doc_splits

    else:
        # Use LLMSherpa for all other file types
        # If exception occurs, raise it

        try:
            # LLMSherpa loader (requires container nlm-ingest to be running)
            # TODO: Include more metadata into each chunk (i.e., which page is the chunk from)
            loader = LLMSherpaFileLoader(
                file_path=file_path,
                new_indent_parser=True,
                apply_ocr=True,
                strategy="text", # this can be "chunks" or "html".
                llmsherpa_api_url="http://nlm-ingestor:5001/api/parseDocument?renderFormat=all",
            )

            st.toast(f"Chunking document {file_name}...")
            docs = loader.load()

            doc_splits = text_splitter.split_documents(docs)

            return doc_splits
            
        except Exception as e:
            st.write(f"Error: {e}")

# Upload files
def upload_files(uploaded_files, st, tenant_id, username=config["neo4j_username"], password=config["neo4j_password"]):
    
    combined_doc_splits = []
    for uploaded_file in uploaded_files:
        doc_path = f"documents/{tenant_id}/{uploaded_file.name}"
        st.toast(f"Processing {uploaded_file.name}")
        # check if file exists
        if not os.path.exists(doc_path):
            st.error(f"File {uploaded_file.name} does not exist in {doc_path}.")
            continue
        doc_splits = docParser(doc_path, st)
        print(f"Number of splits: {len(doc_splits)}")
        print("\n")
        combined_doc_splits.extend(doc_splits)

    print("Writing to database in index:", tenant_id)
    try:
        # stores the parsed documents in the Neo4j database
        Neo4jVector.from_documents(
            documents=combined_doc_splits,
            url=config["neo4j_url"],
            username=username,
            password=password,
            embedding=embeddings,
            index_name=tenant_id,
            node_label=tenant_id,
        )
    except Exception as e:
        st.error(f"Error: {e}")

    print("Documents written to database")

