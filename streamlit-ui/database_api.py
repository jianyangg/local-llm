import os
from docx import Document as docxDocument
from langchain.docstore.document import Document as LangchainDocument
from langchain_community.document_loaders.llmsherpa import LLMSherpaFileLoader
from langchain_community.vectorstores import Neo4jVector
from langchain_community.embeddings import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

config={"ollama_base_url": "http://localhost:11434",
        "llm_name": "llama3",
        "neo4j_url": "bolt://localhost:7687",
        "neo4j_username": "neo4j",
        "neo4j_password": "password",		
        }

# load embedding model
embeddings = OllamaEmbeddings(
    base_url=config["ollama_base_url"],	
    model=config["llm_name"]
)

def docParser(file_path):
    """
    Parses a document file and returns a split up version of the document.
    Requires file type to be reflected in the file extension.

    Args:
        file_path (str): The path to the document file

    Returns:
        List of LangchainDocument objects
    """

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=750, chunk_overlap=20)

    # This assumes that we can tell file type from file extension
    # May not work for linux based systems
    # TODO: NOT THE BEST CHUNKING STRATEGY
    if file_path.endswith('.docx'):
        doc_splits = []
        doc = docxDocument(file_path)
        para_num = 0
        for para in doc.paragraphs:
            if not para.text:
                continue

            langchain_doc_splits = LangchainDocument(
                page_content=para.text,
                metadata={
                    "source": file_path,
                    "chunk_number": para_num,
                    "chunk_type": "para"
                }
            )
            doc_splits.append(langchain_doc_splits)
            para_num += 1

        return doc_splits

    else:
        # Use LLMSherpa for all other file types
        # If exception occurs, raise it

        try:
            # LLMSherpa loader (requires container nlm-ingest to be running)
            # nlm-ingest is port forward to 5010
            loader = LLMSherpaFileLoader(
                file_path=file_path,
                new_indent_parser=True,
                apply_ocr=True,
                strategy="text", # this can be "chunks" or "html".
                llmsherpa_api_url="http://localhost:5010/api/parseDocument?renderFormat=all",
            )

            docs = loader.load()

            doc_splits = text_splitter.split_documents(docs)

            return doc_splits
            
        except Exception as e:
            print(f"Error: {e}")

# Upload files
def upload_files(uploaded_files):
    combined_doc_splits = []
    for uploaded_file in uploaded_files:
        # documents folder should exist
        doc_path = f"documents/{uploaded_file.name}"
        print(f"Processing {doc_path}")
        doc_splits = docParser(doc_path)
        print(f"Number of splits: {len(doc_splits)}")
        print("\n")
        combined_doc_splits.extend(doc_splits)

    print("Writing to database")
    # stores the parsed documents in the Neo4j database
    Neo4jVector.from_documents(
        documents=combined_doc_splits,
        url=config["neo4j_url"],
        username=config["neo4j_username"],
        password=config["neo4j_password"],
        embedding=embeddings,
        index_name="parsers_trial_2",
        node_label="parsersTrial2",
        pre_delete_collection=True,
    )

    print("Documents written to database")

